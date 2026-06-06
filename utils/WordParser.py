import base64
import json
import mimetypes
import os
import uuid
from datetime import datetime
from PIL import Image
from docx import Document as Docx
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import Table
from docx.text.paragraph import Paragraph
from openai import OpenAI
from qwen_token_counter import get_token_count
from models import Document
from utils.ai_endpoint import get_ai_base_url
from utils.error_codes import BizCode

"""
word解析器，使用python-docx提取docx中的文本图像和表格
无法操作doc文件（太老了）
"""

class WordParser:
    def __init__(self):
        self.document_base_dir = os.getenv("DOCUMENT_BASE_DIR", "D:/Pycharm/code/Maintenance_Assistance_System")
        self.document_dir = os.getenv("DOCUMENT_DIR", "upload/documents")
        self.image_dir = os.getenv("IMAGE_DIR", "upload/images")
        self.ai = os.getenv("SERVER_IP", "192.168.246.200")
        self.api_key = os.getenv("API_KEY", "EMPTY")
        self.model = os.getenv("MODEL_AI", "/models/Qwen3-VL-8B-Instruct")
        self.max_token = int(os.getenv("SUMMARY_MAX_TOKEN", 2000))
        self.input_token = int(os.getenv("INPUT_TOKEN", 8000))
        self.last_error_code = None
        self.last_error_detail = None
        base_url = os.path.join(self.document_base_dir, self.document_dir)
        if not os.path.exists(base_url):
            os.makedirs(base_url)
            print(f"创建目录: {base_url}")
        base_url = os.path.join(self.document_base_dir, self.image_dir)
        if not os.path.exists(base_url):
            os.makedirs(base_url)
            print(f"创建目录: {base_url}")

    def get_content(self, file_path):
        if not os.path.exists(file_path):
            raise FileNotFoundError()

        doc = Docx(file_path)
        text_parts = []
        image_urls = []
        file_names = []
        section_image_indexes = self._empty_section_image_indexes()
        current_section = None
        table_index = 0

        for block in self._iter_block_items(doc):
            if isinstance(block, Paragraph):
                block_text = block.text.strip()
                if block_text:
                    detected_section = self._detect_section(block_text)
                    if detected_section:
                        current_section = detected_section
                    text_parts.append(block_text)
                image_indexes = self._save_images_from_element(doc, block._element, image_urls, file_names)
            else:
                table_index += 1
                table_text = f"\n表格{table_index}\n{self.table_to_markdown(block)}"
                detected_section = self._detect_section(table_text)
                if detected_section:
                    current_section = detected_section
                text_parts.append(table_text)
                image_indexes = self._save_images_from_element(doc, block._element, image_urls, file_names)

            for image_index in image_indexes:
                text_parts.append(f"【图片{image_index}】")
                if current_section:
                    section_image_indexes[current_section].append(image_index)

        return "\n".join(text_parts), image_urls, file_names, section_image_indexes

    def _iter_block_items(self, doc):
        for child in doc.element.body.iterchildren():
            if isinstance(child, CT_P):
                yield Paragraph(child, doc)
            elif isinstance(child, CT_Tbl):
                yield Table(child, doc)

    def _extract_image_rel_ids(self, element):
        rel_ids = []
        for xpath in (".//a:blip/@r:embed", ".//a:blip/@r:link"):
            try:
                rel_ids.extend(element.xpath(xpath))
            except Exception:
                continue
        return rel_ids

    def _save_image_part(self, image_part):
        img_ext = image_part.content_type.split('/')[-1]
        if img_ext == "jpeg":
            img_ext = "jpg"

        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        unique_filename = f"{timestamp}_{uuid.uuid4().hex}.{img_ext}"
        base_url = os.path.join(self.document_base_dir, self.image_dir)
        img_path = os.path.join(base_url, unique_filename)
        with open(img_path, 'wb') as f:
            f.write(image_part.blob)
        print(f"图片已保存: {img_path}")
        return img_path, unique_filename

    def _save_images_from_element(self, doc, element, image_urls, file_names):
        image_indexes = []
        for rel_id in self._extract_image_rel_ids(element):
            if rel_id not in doc.part.rels:
                continue
            rel = doc.part.rels[rel_id]
            if "image" not in rel.target_ref:
                continue
            img_path, unique_filename = self._save_image_part(rel.target_part)
            image_urls.append(img_path)
            file_names.append(unique_filename)
            image_indexes.append(len(image_urls))
        return image_indexes

    def _detect_section(self, text: str):
        normalized = "".join((text or "").split())
        section_keywords = [
            ("problem_intro", ["问题描述", "问题简介"]),
            ("causes", ["原因分析", "原因"]),
            ("evaluation", ["故障评估", "评估"]),
            ("inspection", ["检查步骤", "检查"]),
            ("solutions", ["解决方案", "问题解决", "解决"]),
            ("key_points", ["关键要点", "总结"]),
        ]
        for section, keywords in section_keywords:
            if any(keyword in normalized for keyword in keywords):
                return section
        return None

    def _empty_section_image_indexes(self):
        return {
            "problem_intro": [],
            "causes": [],
            "evaluation": [],
            "inspection": [],
            "solutions": [],
            "key_points": [],
        }
    def _set_last_error(self, code: int, message: str):
        self.last_error_code = int(code)
        self.last_error_detail = message

    def _is_ai_service_unavailable_error(self, error: Exception) -> bool:
        name = type(error).__name__
        if name in {"APIConnectionError", "APITimeoutError"}:
            return True
        msg = str(error).lower()
        keywords = [
            "connection",
            "timed out",
            "timeout",
            "refused",
            "temporarily unavailable",
            "service unavailable",
            "name resolution",
            "max retries exceeded",
            "502",
            "503",
            "504",
        ]
        return any(k in msg for k in keywords)

    def table_to_markdown(self, table):
        """
        将表格转换为Markdown表格字符串。
        """
        # 获取表格的行数和列数（考虑合并单元格）
        rows = table.rows
        if not rows:
            return ""

        # 先确定列数：取第一行单元格数，但可能因合并而少于实际列数
        # 更稳健的方式：遍历所有单元格，找到最大列索引
        col_count = max(len(row.cells) for row in rows)

        # 构建矩阵，初始填充空字符串
        matrix = [["" for _ in range(col_count)] for _ in range(len(rows))]

        # 填充单元格内容
        for i, row in enumerate(rows):
            for j, cell in enumerate(row.cells):
                if j < col_count:  # 防止索引越界
                    # 如果单元格已合并，其左上角位置才是有效内容，其余位置应保持空
                    # 简单处理：每个单元格只填一次，合并的单元格只保留第一个
                    matrix[i][j] = cell.text.strip().replace('\n', ' ')  # 换行符替换为空格

        # 生成 Markdown 表格
        md_lines = []
        # 表头
        header = matrix[0]
        md_lines.append("| " + " | ".join(header) + " |")
        # 分隔线
        md_lines.append("|" + "|".join([" --- " for _ in range(col_count)]) + "|")
        # 数据行
        for row in matrix[1:]:
            md_lines.append("| " + " | ".join(row) + " |")

        return "\n".join(md_lines)

    def image_to_base64(self, image: str):
        with open(image, "rb") as f:
            image_base64 = base64.b64encode(f.read()).decode("utf-8")
            return image_base64

    def compress_image(self, image_path: str, max_size=512, pad_color=(0, 0, 0)):
        if not os.path.exists(image_path):
            raise FileNotFoundError()

        image = Image.open(image_path).convert("RGB")
        max_length = max(image.width, image.height)
        rate = max_size / max_length
        new_size = (int(image.width * rate), int(image.height * rate))
        resized_image = image.resize(new_size)

        new_image = Image.new("RGB", (max_size, max_size), pad_color)

        x = (max_size - new_size[0]) // 2
        y = (max_size - new_size[1]) // 2

        new_image.paste(resized_image, (x, y))

        dir_name, filename = os.path.split(image_path)
        name, ext = os.path.splitext(filename)
        new_path = f"{name}_compressed_{max_size}{ext}"
        new_path = os.path.join(dir_name, new_path)
        new_image.save(new_path)
        return new_path

    def parse(self, file_path):
        self.last_error_code = None
        self.last_error_detail = None
        text, image_urls, file_names, section_image_indexes = self.get_content(file_path)
        document = self.file2Document(text, image_urls, file_names, section_image_indexes)
        return document

    def generate_message(self, text: str = None, image_urls: str = None):
        messages = []
        data = {}
        # msg_content = []
        prompt = """你好，你是一位问题分析专家，我将给你一段有关设备维修的文本和几张图片。请你最大限度使用内容，按照以下的模板提供信息，以JSON格式返回。若内容未提供，字段可以为空，但不可缺失字段。对于图片，请给出图片编号。
模板：
标题：<简洁的标题，点明核心内容>
问题简介：<可包含定义解释，现象介绍，问题发生频率，后果等内容>
原因：<造成该问题的主要原因，尽量从高频到低频排序>
评估：<评估问题的手段，方法，工具等信息>
检查：<描述维修现场如何进行定位确认>
解决方法：<现场的解决措施及根本的解决方案>
总结：<总结问题的主要原因，后果及解决方案的关键信息>

输出JSON格式如下：
{{
    "title": "标题", // 案例名
    "problem_intro": "问题简介文本",
    "image_urls_problem_intro": [1, 3], // 与问题简介相关的图片编号
    "causes": "原因文本",
    "image_urls_causes": [2], // 与原因相关的图片编号
    "evaluation": "评估方法文本",
    "image_urls_evaluation": [4], // 与评估相关的图片编号
    "inspection": "检查方法文本",
    "image_urls_inspection": [5], // 与检查相关的图片编号
    "solutions": "解决方案文本",
    "image_urls_solutions": [6], // 与解决方案相关的图片编号
    "key_points": "总结文本",
    "image_urls_key_points": [7, 8] // 与总结相关的图片编号
}}

注意：
1.所有内容必须严格来源于提供的文本和图片。禁止任何形式的脑补、推理、常识补充或添加原文未提及的修饰词与解释，并且最大限度利用文本。
2. 内容中不包含的信息，对应字段可以为空，若不包含图片，图片为空列表[]。
3. 所有文字必须100%来自文档原文和图片，不可杜撰任何信息。
4. 你给出的回答仅包含我要求的JSON格式答案。
5. 给定图片编号从1开始，不得编写新的图片。
6. title不可为空，同一张图片不要出现太多次，即一张图片不要出现在2个以上字段中。
7. 内容需连贯详细，最大限度使用给定内容，请勿过分精简。
8. 检查步骤和解决方案字段若有内容相关，请尽可能详细描述。
9. 各字段内容请勿大量重复。

现在请分析下面的内容：
[文本内容]
{text}

[图片内容由后续base64给出]
""".format(text=text)
        prompt += "\n注意：图片归属必须优先依据其在原文中的位置。图片位于哪个小节下，就归入对应 image_urls 字段，不得仅凭图片内容语义移动到其他字段。"
        # prompt = f""
        msg_content = [{"type": "text", "text": prompt}]
        # encoding = tiktoken.get_encoding("cl100k_base")
        token_cnt = get_token_count(prompt)

        print(f"token1: {token_cnt}")
        for image in image_urls:
            print(image)
            if token_cnt >= self.input_token - 1000:
                break
            compress_image = self.compress_image(image)
            mime_type, _ = mimetypes.guess_type(compress_image)
            if mime_type is None:
                ext = os.path.splitext(compress_image)[1].lower()
                mime_type = {
                    '.png': 'image/png',
                    '.jpg': 'image/jpeg',
                    '.jpeg': 'image/jpeg',
                    '.webp': 'image/webp',
                    '.bmp': 'image/bmp'
                }.get(ext, 'image/jpeg')
            image_base64 = self.image_to_base64(compress_image)
            msg_content.append({
                "type": "image_url",
                "image_url": {"url": f"data:{mime_type};base64,{image_base64}"}
            })
            token_cnt += 258
        data["role"] = "user"
        data["content"] = msg_content
        messages.append(data)
        print(f"tokens: {token_cnt}")
        return messages

    def _image_indexes_to_urls(self, image_indexes, image_names):
        urls = []
        for image_index in image_indexes:
            if image_index <= 0 or image_index > len(image_names):
                continue
            urls.append(self.image_dir + "/" + image_names[image_index - 1])
        return ", ".join(urls) if urls else None

    def _apply_section_image_urls(self, result, section_image_indexes, image_names):
        if not section_image_indexes or not any(section_image_indexes.values()):
            return result, set()

        used_indexes = set()
        for section, image_indexes in section_image_indexes.items():
            field = f"image_urls_{section}"
            result[field] = self._image_indexes_to_urls(image_indexes, image_names)
            used_indexes.update(image_indexes)
        return result, used_indexes

    def file2Document(self, text, image_urls, image_names, section_image_indexes=None):
        try:
            client = OpenAI(
                base_url=get_ai_base_url(),
                api_key=self.api_key
            )

            messages = self.generate_message(text, image_urls)

            response = client.chat.completions.create(
                model=self.model,
                messages=messages,
                max_tokens=self.max_token
            )
            print(response)
            ans = response.choices[0].message.content
            print(ans)
            result = json.loads(ans)

            result, used_image_indexes = self._apply_section_image_urls(result, section_image_indexes, image_names)

            if not used_image_indexes:
                flag = [0] * len(image_names)
                for key in result.keys():
                    if "image" in key:
                        image_url_content = ""
                        for image_index in result[key]:
                            if image_index > len(image_urls) or flag[image_index - 1] == 1:
                                continue
                            url = image_names[image_index - 1]
                            flag[image_index - 1] = 1
                            url = self.image_dir + "/" + url
                            image_url_content += url + ", "
                        image_url_content = image_url_content.rstrip(", ")
                        if len(result[key]) == 0:
                            image_url_content = None
                        result[key] = image_url_content
                used_image_indexes = {i + 1 for i, used in enumerate(flag) if used == 1}
            document = Document(**result,
                                is_vectorized=0)

            if len(image_urls) > 0:
                for i in range(len(image_urls)):
                    if i + 1 not in used_image_indexes:
                        os.remove(image_urls[i])
                        dir_name, filename = os.path.split(image_urls[i])
                        name, ext = os.path.splitext(filename)
                        new_path = f"{name}_compressed.{ext}"
                        new_path = os.path.join(dir_name, new_path)
                        if os.path.exists(new_path):
                            os.remove(new_path)
            # print(type(document))
            return document

        except Exception as e:
            print(e)
            if self._is_ai_service_unavailable_error(e):
                self._set_last_error(BizCode.AI_SERVICE_UNAVAILABLE, "AI服务不可用，请稍后重试")
            else:
                self._set_last_error(BizCode.DOC_PARSE_FAILED, str(e))
            for image in image_urls:
                if os.path.exists(image):
                    os.remove(image)
            return None

word_parser = WordParser()

if __name__ == "__main__":
    pass
